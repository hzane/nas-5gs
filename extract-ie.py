#!/usr/bin/env python
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table, _Row
from docx.text.paragraph import Paragraph


def block_items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    elif isinstance(parent, _Row):
        parent_elm = parent._tr
    else:
        raise ValueError("something's not right")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

fn = '24501-g10.docx'
doc = Document(fn)

for block in block_items(doc):
    if isinstance(block, Paragraph):
        last_para = block
    elif isinstance(block, Table):
        cols = []
        for cell in block.rows[0].cells:
            txts = []
            for p in cell.paragraphs:
                txts.append(p.text)
            col = ','.join(txts)
            cols.append(col)
        if cols[0] == 'IEI' and last_para.text.startswith('Table'):
            print(last_para.text)
            print('\t'.join(cols))

            for row in block.rows[1:]:
                line = []
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if not p.text.startswith('NOTE:'):
                            line.append(p.text)
                print('\t'.join(line))
            print()
